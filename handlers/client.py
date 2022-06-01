from aiogram import types, Dispatcher
from creation import dp, bot
from keybuttons.keybuttons import *
from inline import *
from aiogram.types import ReplyKeyboardRemove
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
import time, re, random
from datetime import datetime
import docx
import os


class Fsm_storage(StatesGroup):
    # login = State()
    fio = State()
    day_of_birth = State()
    place_of_birth = State()
    birth_time = State()
    source_of_info = State()
    circumstances = State()
    marriage = State()
    divorce = State()
    children = State()
    relocation = State()
    job = State()
    events = State()
    death = State()
    operations = State()
    property = State()
    other = State()
    wishes = State()
    contacts = State()
    adds = State()

async def main(message: types.Message):
    await bot.send_message(message.from_user.id, 'Добрый день! Меня зовут Алла Бабенко и я Профессиональный Астролог.\n'
                                                 'Для того, чтобы построить Вашу Натальную карту или составить '
                                                 'Прогностику, мне необходимо знать Вашу Дату рождения, Место рождения '
                                                 'и точное Время рождения.\nЕсли время рождения неизвестно или есть '
                                                 'сомнения в его точности, значит необходимо рассчитать Ректификацию '
                                                 'времени рождения (уточнение времени рождения через астрологические '
                                                 'методы вычисления). \nДля этого мне понадобится более подробная '
                                                 'информация о Вас. \nПоэтому прошу Вас ответить на следующие вопросы.')

    await bot.send_message(message.from_user.id, 'Если Вы готовы заполнить Астрологическую анкету, нажмите пожалуйста «Начать заполнять Анкету».',
                           reply_markup=begin_filling)




async def yes_begin_filling(message: types.Message, state: FSMContext):
    # await state.reset()
    await bot.send_message(message.from_user.id, 'Напишите пожалуйста Ваше ФИО полностью.')
    await Fsm_storage.fio.set()


async def no_begin_filling(message: types.Message):
    await bot.send_message(message.from_user.id, 'Хорошо! Как будете готовы, нажмите /start')


async def fio_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['fio'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Напишите пожалуйста Вашу Дату рождения в формате чч.мм.гггг.')
        await Fsm_storage.day_of_birth.set()



async def day_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:

        data['day_of_birth'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Напишите пожалуйста Ваше Место рождения полностью '
                                                     f'(страна, область/регион, город/населённый пункт).')
        await Fsm_storage.place_of_birth.set()


async def place_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['place_of_birth'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Напишите пожалуйста Ваше Время рождения с указанием части дня '
                                                     f'(например - 5-30 утра). \nЕсли Время рождения неизвестно, укажите '
                                                     f'приблизительный временной диапазон.')
        await Fsm_storage.birth_time.set()


async def birth_time_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['birth_time'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста источник информации о Времени рождения '
                                                     f'(бирка/со слов мамы или др.)')
        await Fsm_storage.source_of_info.set()


async def source_of_info_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['source_of_info'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Какими были обстоятельства Вашего рождения '
                                                     f'(естественные роды/кесарево, запланировано/не запланировано, '
                                                     f'в больнице/нет, быстро/долго, стимуляция, необычные факты).')
        await Fsm_storage.circumstances.set()


async def circumstances_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['circumstances'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Далее необходимо указать события из Вашей жизни для того, '
                                                     f'чтобы провести расчёт точного Времени Вашего рождения.')

        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Вы когда-нибудь состояли в браке?', reply_markup=yes7)



async def yeshandler7(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию относительно Ваших браков '
                                                     f'\n(необходимо указать все браки, которые были/есть в Вашей жизни) '
                                                     f'в следующем формате: \n\nБрак 1-й, дата чч.мм.гггг. И '
                                                     f'комментарии по событиям этого брака.')
        await Fsm_storage.marriage.set()


async def nohandler7(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['marriage'] = 'В браке не состоял'
        data['divorce'] = 'Разводов не было'
        await state.finish()
        await bot.send_message(message.from_user.id, f'Рождение детей (необходимо указать всех рождённых Вами детей).\n'
                                                     f'Напишите пожалуйста полную Дату рождения Вашего ребёнка и '
                                                     f'комментарии относительно Рождения ребёнка в следующем '
                                                     f'формате:\n\n'
                                                     f'Ребёнок 1-й дата рождения чч.мм.гггг, комментарий относительно '
                                                     f'Рождения.')
        await Fsm_storage.children.set()

async def marriage_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['marriage'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Если был развод по вышеуказанному браку/бракам, укажите '
                                                     'пожалуйста полную дату чч.мм.гггг, когда был осуществлён '
                                                     'бракоразводный процесс и комментарии по деталям этого развода. \n'
                                                     'Если развода не было, укажите что этого события в Вашей '
                                                     'жизни не случалось. Информацию необходимо указывать в '
                                                     'следующем формате: \n\nРазвод 1-й дата, указанная в документе '
                                                     'чч.мм.гггг. И комментарии по событиям этого развода.')
        await Fsm_storage.divorce.set()




async def divorce_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['divorce'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Рождение детей (необходимо указать всех рождённых Вами детей). '
                                                     f'Напишите пожалуйста полную Дату рождения Вашего ребёнка и '
                                                     f'комментарии относительно Рождения ребёнка в следующем формате:\n\n '
                                                     f'Ребёнок 1-й дата рождения чч.мм.гггг, комментарий относительно '
                                                     f'Рождения.')
        await Fsm_storage.children.set()




async def children_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['children'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Укажите пожалуйста информацию о переездах, случавшихся в '
                                                     'Вашей жизни (из города в город, из страны в страну), в '
                                                     'следующем формате:\n\nПереезд из ____ в ____   дата переезда '
                                                     'чч.мм.гггг, комментарий относительно этого переезда '
                                                     '(вынужденный, по работе, в связи с какими обстоятельствами, '
                                                     'как сказался на Вашей жизни).')

        await Fsm_storage.relocation.set()






async def relocation_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:

        data['relocation'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Укажите пожалуйста информацию обо всех Ваших местах работы '
                                                     '(приём/увольнение) в следующем формате:\n\nРабота 1-я дата приёма '
                                                     '(чч.мм.гггг), дата увольнения (чч.мм.гггг), комментарий '
                                                     'относительно этого места работы')

        await Fsm_storage.job.set()





async def job_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['job'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Укажите пожалуйста были ли в Вашей жизни события, в '
                                                     'результате которых Вы получали награды, достигали побед '
                                                     'в различных сферах, либо просто значимые достижения\n'
                                                     '(необходимо написать все самые выдающееся события) в '
                                                     'следующем формате:\n\nСобытие 1-е дата получения чч.мм.гггг, '
                                                     'комментарий относительно этого достижения.')

        await Fsm_storage.events.set()





async def events_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['events'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               'Укажите пожалуйста информацию относительно смерти родственников/близких людей, '
                               'в следующем формате:\n\nСобытие 1-е дата происшедшего чч.мм.гггг, '
                               'комментарий относительно этой даты.')
        await Fsm_storage.death.set()





async def death_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['death'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию относительно травм/операций в '
                                                     f'Вашей жизни, в следующем формате:\n\nСобытие 1-е дата '
                                                     f'случившегося чч.мм.гггг, '
                                                     f'комментарий относительно этого происшествия.')
        await Fsm_storage.operations.set()





async def operations_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:

        data['operations'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               f'Укажите пожалуйста информацию о покупке движимого и недвижимого имущества '
                               f'(автомобиль, квартира и т.д.), в следующем формате:\n\nСобытие 1-е дата приобретения '
                               f'чч.мм.гггг, комментарий относительно этой покупки.')
        await Fsm_storage.property.set()




async def property_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['property'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               f'Если в Вашей жизни происходило иное значимое событие, которое не упоминалось в '
                               f'вышеуказанном списке вопросов, просьба указать его в следующем формате:\n\nСобытие '
                               f'1-е и его дата чч.мм.гггг, комментарий относительно происшедшего. Если добавить '
                               f'больше нечего, так и напишите.')
        await Fsm_storage.other.set()



async def other_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['other'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               f'Напишите пожалуйста, что Вас интересует в разборе Натальной карты/Прогностике? '
                               f'И на чём следует акцентировать внимание?\n\nТакже можете просто перечислить сферы '
                               f'жизни, которые хотелось бы обсудить во время консультации (личная жизнь, '
                               f'финансы, работа и пр.).')
        await Fsm_storage.wishes.set()



async def wishes_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['wishes'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Укажите пожалуйста Ваши контактные данные – номер телефона '
                                                     'и адрес электронной почты.')

        await Fsm_storage.contacts.set()


async def precheck(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['contacts'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Проверьте пожалуйста правильно ли Вы указали данные по '
                                                     'вышеперечисленным вопросам, и если все верно, то нажмите '
                                                     '"Завершить и отправить анкету". Если хотите что-то исправить, '
                                                     'то нажмите "Дополнить"', reply_markup=redact)


async def redaction(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, 'Напишите пожалуйста информацию, для того, чтобы скорректировать '
                                                     'или дополнить ваши ответы.')

        await Fsm_storage.adds.set()

async def check_adds(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['adds'] = message.text
        await state.finish()

        await bot.send_message(message.from_user.id,
                               'Благодарю Вас за время, которое Вы уделили, чтобы ответить на все вышеперечисленные '
                               'вопросы. \nХорошего вам дня! И До встречи!')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               'Контактные данные для связи:\nАлла Бабенко\nFreyaforseti0609@mail.ru '
                               '\nFreyaforseti0609@gmail.com\n+79998482603')
        time.sleep(2)
        await bot.send_message(message.from_user.id,
                               'Так же вы можете воспользоваться QR-кодом для того, чтобы сохранить контактный номер телефона.')
        qr = open("static/qr.jpg", 'rb')
        await bot.send_photo(message.from_user.id, photo=qr)

        # login = f'@{message.from_user.username}'
        data_list = [data['fio'], data['day_of_birth'], data['place_of_birth'], data['birth_time'], data['source_of_info'],
                     data['circumstances'], data['marriage'], data['divorce'], data['children'], data['relocation'],
                     data['job'], data['events'], data['death'], data['operations'], data['property'],
                     data['other'], data['wishes'], data['contacts'], data['adds']]

        list_of_points = ['ФИО', 'Дата рождения', 'Место рождения', 'Время рождения', 'Источник информации о времени рождения',
                          'Обстоятельства рождения', 'Брак', 'Развод', 'Дети', 'Переезд', 'Работа',
                          'Важные события в жизни', 'Смерть родственников', 'Травмы и операции',
                          'Частная собственность', 'Прочие события', 'Уточнение по Консультации', 'Контакты для связи',
                          'Дополнения к ответам']
        mydoc = docx.Document()

        mydoc.add_heading(f"Анкета клиента {data['fio']}", 1)
        mydoc.add_paragraph('')
        for i in range(len(list_of_points)):
            p = mydoc.add_paragraph('')
            p.add_run(f'{list_of_points[i]}: ').bold = True
            p.add_run(f'{data_list[i]}')

        mydoc.save(f"{data['fio']}.docx")
        file = open(f"{data['fio']}.docx", 'rb')
        await bot.send_document(363700041, document=file)
        await bot.send_document(277962812, document=file)
        try:
            await bot.send_message(363700041, f'Пользователь: @{message.from_user.username}')
            await bot.send_message(277962812, f'Пользователь: @{message.from_user.username}')
        except:
            await bot.send_message(363700041, f'У пользователя нет username. Необходимо проверить "Контакты для связи" '
                                              f'в файле выше.')
            await bot.send_message(277962812, f'У пользователя нет username. Необходимо проверить "Контакты для связи" '
                                              f'в файле выше.')
        os.remove(f"{data['fio']}.docx")
    await state.finish()
    await state.reset_state()


async def final(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id,
                               'Благодарю Вас за время, которое Вы уделили, чтобы ответить на все вышеперечисленные '
                               'вопросы. Хорошего вам дня! И До встречи!')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               'Контактные данные для связи:\nАлла Бабенко\nFreyaforseti0609@mail.ru '
                               '\nFreyaforseti0609@gmail.com\n+79998482603')
        time.sleep(2)
        await bot.send_message(message.from_user.id,
                               'Так же вы можете воспользоваться QR-кодом для того, чтобы сохранить контактный номер телефона.')
        qr = open("static/qr.jpg", 'rb')
        await bot.send_photo(message.from_user.id, photo=qr)


        # login = f'@{message.from_user.username}'
        data_list = [data['fio'], data['day_of_birth'], data['place_of_birth'], data['birth_time'], data['source_of_info'],
                     data['circumstances'], data['marriage'], data['divorce'], data['children'], data['relocation'],
                     data['job'], data['events'], data['death'], data['operations'], data['property'],
                     data['other'], data['wishes'], data['contacts']]

        list_of_points = ['ФИО', 'Дата рождения', 'Место рождения', 'Время рождения', 'Источник информации о времени рождения',
                          'Обстоятельства рождения', 'Брак', 'Развод', 'Дети', 'Переезд', 'Работа',
                          'Важные события в жизни', 'Смерть родственников', 'Травмы и операции',
                          'Частная собственность', 'Прочие события', 'Уточнение по Консультации', 'Контакты для связи']
        mydoc = docx.Document()

        mydoc.add_heading(f"Анкета клиента {data['fio']}", 1)
        mydoc.add_paragraph('')
        for i in range(len(list_of_points)):
            p = mydoc.add_paragraph('')
            p.add_run(f'{list_of_points[i]}: ').bold = True
            p.add_run(f'{data_list[i]}')

        mydoc.save(f"{data['fio']}.docx")
        file = open(f"{data['fio']}.docx", 'rb')
        await bot.send_document(363700041, document=file)
        await bot.send_document(277962812, document=file)
        try:
            await bot.send_message(363700041, f'Пользователь: @{message.from_user.username}')
            await bot.send_message(277962812, f'Пользователь: @{message.from_user.username}')
        except:
            await bot.send_message(363700041, f'У пользователя нет username. Необходимо проверить "Контакты для связи" '
                                              f'в файле выше.')
            await bot.send_message(277962812, f'У пользователя нет username. Необходимо проверить "Контакты для связи" '
                                              f'в файле выше.')
        os.remove(f"{data['fio']}.docx")
    await state.finish()
    await state.reset_state()


async def stats(message: types.Message):
    print(message)
    print(message.from_user)
    await bot.send_message(363700041, f'{message.from_user.id}')
    await bot.send_message(363700041, f'Пользователь: {message.from_user.username}, id: {message.from_user.id}')


async def lost(message: types.Message):
    await bot.send_message(message.from_user.id, 'Не понимаю Вас, нажмите --- /start')


def register_handlers_client(dp: Dispatcher):
    dp.register_message_handler(main, commands=['start', 'help'])
    dp.register_callback_query_handler(yes_begin_filling, text='now')
    dp.register_callback_query_handler(no_begin_filling, text='later')
    dp.register_message_handler(fio_check, state=Fsm_storage.fio)
    dp.register_message_handler(day_of_birth_check, state=Fsm_storage.day_of_birth)
    dp.register_message_handler(birth_time_check, state=Fsm_storage.birth_time)
    dp.register_message_handler(place_of_birth_check, state=Fsm_storage.place_of_birth)
    dp.register_message_handler(source_of_info_check, state=Fsm_storage.source_of_info)
    dp.register_message_handler(circumstances_check, state=Fsm_storage.circumstances)
    dp.register_callback_query_handler(yeshandler7, text='yes7')
    dp.register_callback_query_handler(nohandler7, text='no7')
    dp.register_message_handler(marriage_check, state=Fsm_storage.marriage)
    dp.register_message_handler(divorce_check, state=Fsm_storage.divorce)
    dp.register_message_handler(children_check, state=Fsm_storage.children)
    dp.register_message_handler(relocation_check, state=Fsm_storage.relocation)
    dp.register_message_handler(job_check, state=Fsm_storage.job)
    dp.register_message_handler(events_check, state=Fsm_storage.events)
    dp.register_message_handler(death_check, state=Fsm_storage.death)
    dp.register_message_handler(operations_check, state=Fsm_storage.operations)
    dp.register_message_handler(property_check, state=Fsm_storage.property)
    dp.register_message_handler(other_check, state=Fsm_storage.other)
    dp.register_message_handler(wishes_check, state=Fsm_storage.wishes)
    dp.register_message_handler(precheck, state=Fsm_storage.contacts)
    dp.register_callback_query_handler(redaction, text='redaction')
    dp.register_message_handler(check_adds, state=Fsm_storage.adds)
    dp.register_callback_query_handler(final, text='ready')
    # dp.register_message_handler(final, state=Fsm_storage.contacts)
    dp.register_message_handler(stats, commands=['stats'])
    dp.register_message_handler(lost)

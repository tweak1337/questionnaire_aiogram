from aiogram import types, Dispatcher
from creation import dp, bot
from keybuttons.keybuttons import *
from inline import *
from aiogram.types import ReplyKeyboardRemove
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
import time, re, random
from datetime import datetime
import docx
import os


class Fsm_storage(StatesGroup):
    # login = State()
    fio = State()
    year_of_birth = State()
    month_of_birth = State()
    day_of_birth = State()
    place_of_birth = State()
    birth_time = State()
    source_of_info = State()
    circumstances = State()
    not_in_hospital = State()
    childbirth = State()
    marriage = State()
    divorce = State()
    children = State()
    relocation = State()
    job = State()
    events = State()
    death = State()
    operations = State()
    property = State()
    other = State()
    wishes = State()
    contacts = State()

async def main(message: types.Message):
    await bot.send_message(message.from_user.id, 'Добрый день! Меня зовут Алла Бабенко и я Профессиональный Астролог. '
                                                 'Для того, чтобы построить Вашу Натальную карту или составить '
                                                 'Прогностику, мне необходимо знать Вашу Дату рождения, Место рождения '
                                                 'и точное Время рождения.Если время рождения неизвестно или есть '
                                                 'сомнения в его точности, значит необходимо рассчитать Ректификацию '
                                                 'времени рождения (уточнение времени рождения через астрологические '
                                                 'методы вычисления). Для этого мне понадобится более подробная '
                                                 'информация о Вас. Поэтому прошу Вас ответить на следующие вопросы.')

    await bot.send_message(message.from_user.id, 'Нажмите /Готов, если готов, и /Не_готов, если пока не готов',
                           reply_markup=kb_gotov)


async def gotov(message: types.Message, state: FSMContext):
    # await state.reset()
    await bot.send_message(message.from_user.id, 'Напишите пожалуйста Ваше ФИО полностью.',
                           reply_markup=ReplyKeyboardRemove())
    await Fsm_storage.fio.set()


async def negotov(message: types.Message):
    await bot.send_message(message.from_user.id, 'Хорошо! Как будете готовы, нажми /start',
                           reply_markup=ReplyKeyboardRemove())


async def fio_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['fio'] = message.text
        # data['login'] = message.from_user.username
        await bot.send_message(message.from_user.id, f'Ваше ФИО "{data["fio"]}"?', reply_markup=yes1)

        await state.finish()


async def yeshandler1(message: types.CallbackQuery):
    await bot.send_message(message.from_user.id, f'Отлично! Теперь необходиимо знать дату рождения, год, месяц, '
                                                 f'число и время. Если вдруг не знаете время, можно узнать и '
                                                 f'поставить прочерк, '
                                                 f'но до этого мы еще дойдем.'
                                                 f'Сейчас, пожалуйста укажи сперва год, нужно всего 4 цифры, '
                                                 f'например так: 1995')
    # await state.finish()
    await Fsm_storage.year_of_birth.set()


async def nohandler1(message: types.CallbackQuery):
    await bot.send_message(message.from_user.id, 'Тогда введите корректные данные')
    await Fsm_storage.fio.set()


async def year_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        isdigital = 0
        current_year = datetime.now().year
        for i in message.text:
            if not i.isdigit():
                isdigital += 1
        if isdigital > 0:
            await bot.send_message(message.from_user.id,
                                   'В вашем сообщении должны быть только цифры, пожалуйста, введите еще раз год рождения.')
            await Fsm_storage.year_of_birth.set()
        elif int(message.text) > current_year:
            x = int(message.text) - current_year
            if (x > 14 or x == 1) and int(str(x)[-1]) == 1:
                let = 'год'
            elif (x > 14 or x == 2 or x == 3 or x == 4) and (
                    int(str(x)[-1]) == 2 or int(str(x)[-1]) == 3 or int(str(x)[-1]) == 4):
                let = 'года'
            else:
                let = 'лет'

            await bot.send_message(message.from_user.id,
                                   f'Вероятно, вы прилетели к нам из далекого будущего, тогда вернемся к анкете через {x} {let}')

            await bot.send_message(message.from_user.id,
                                   f'Серьезно, напишите свой настоящий год рождения =) Эти данные нужны для более '
                                   f'точного анализа')

            await Fsm_storage.year_of_birth.set()
        elif current_year - int(message.text) > 100:
            x = current_year - int(message.text)
            if x > 10 and x < 20:
                let = 'лет'
            elif (x > 14 or x == 1) and int(str(x)[-1]) == 1:
                let = 'год'
            elif (x > 14 or x == 2 or x == 3 or x == 4) and (
                    int(str(x)[-1]) == 2 or int(str(x)[-1]) == 3 or int(str(x)[-1]) == 4):
                let = 'года'
            else:
                let = 'лет'
            await bot.send_message(message.from_user.id,
                                   f'Очень сомневаюсь, что вам {x} {let}. '
                                   f'Напишите настоящий год рождения, Эти данные нужны для более точного анализа')
        elif current_year - int(message.text) < 8:
            await bot.send_message(message.from_user.id,
                                   f'Серьезно, напишите свой настоящий год рождения =) Эти данные нужны для '
                                   f'более точного анализа')
            await Fsm_storage.year_of_birth.set()

        else:
            data['year_of_birth'] = message.text
            await state.finish()
            await bot.send_message(message.from_user.id,
                                   f'Принято, теперь напишите, в каком месяце вы родились. Просто числом')
            # await state.finish()
            await Fsm_storage.month_of_birth.set()


async def month_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['month_of_birth'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Отлично! теперь напишите день своего рождения.')
        await Fsm_storage.day_of_birth.set()
        # await state.finish()


async def day_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:

        try:
            digit_check = int(message.text)
        except Exception:
            await bot.send_message(message.from_user.id, f'Пожалуйста, вводите только цифры')
            await Fsm_storage.day_of_birth.set()
            return
        if digit_check > 31 or digit_check == 0:
            await bot.send_message(message.from_user.id,
                                   f'Такой даты не бывает, пожалуйста, введите настоящий день рождения')
            await Fsm_storage.day_of_birth.set()
        else:
            data['day_of_birth'] = message.text
            await state.finish()
            await bot.send_message(message.from_user.id, f'Напишите пожалуйста Ваше Место рождения полностью '
                                                         f'(страна, область/регион, город/населённый пункт).')
            await Fsm_storage.place_of_birth.set()


async def place_of_birth_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['place_of_birth'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Хорошо, двигаемся дальше!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Напишите пожалуйста Ваше Время рождения с указанием части дня '
                                                     f'(например - 5-30 утра). Если Время рождения неизвестно, укажите '
                                                     f'приблизительный временной диапазон.')
        await Fsm_storage.birth_time.set()


async def birth_time_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['birth_time'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста источник информации о Времени рождения '
                                                     f'(бирка/со слов мамы или др.)')
        await Fsm_storage.source_of_info.set()


async def source_of_info_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['source_of_info'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Хорошо, идём дальше!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Какими были обстоятельства Вашего рождения '
                                                     f'(естественные роды/кесарево, запланировано/не запланировано, '
                                                     f'в больнице/нет, быстро/долго, стимуляция, необычные факты).')
        await Fsm_storage.circumstances.set()


async def circumstances_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['circumstances'] += f', {message.text}'
        except Exception as e:
            data['circumstances'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, f'Хотите что-то добавить по обстоятельствам рождения?',
                               reply_markup=yes6)


async def yeshandler6(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Дополните ответным сообщением.')
        await Fsm_storage.circumstances.set()


async def nohandler6(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда пойдем дальше.')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Далее необходимо указать события из Вашей жизни для того, '
                                                     f'чтобы провести расчёт точного Времени Вашего рождения.')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию относительно Ваших браков '
                                                     f'(необходимо указать все браки, которые были/есть в Вашей жизни)')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Скажите, вы когда-нибудь состояли в браке?', reply_markup=yes7)


async def yeshandler7(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, пожалуйста, напишите про первый '
                                                     f'брак (а затем, про последующие, если такие есть) в следующем '
                                                     f'формате: дата (чч.мм.гггг). '
                                                     f'И комментарии по событиям этого брака.')
        await Fsm_storage.marriage.set()


async def nohandler7(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['marriage'] = 'В браке не состоял'
        data['divorce'] = 'Разводов не было'
        await state.finish()
        await bot.send_message(message.from_user.id, f'Хорошо, двигваемся дальше.')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'У вас есть/были дети?', reply_markup=yes11)


async def marriage_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['marriage'] += f', {message.text}'
        except Exception as e:
            data['marriage'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes8)


async def yeshandler8(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Дополните ответным сообщением.')
        await Fsm_storage.marriage.set()


async def nohandler8(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, далее необходимо указать информацию по поводу разводов, '
                                                     f'если такие были.')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Скажите, у вас когда-нибудь были разводы?', reply_markup=yes9)


async def yeshandler9(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста полную дату чч.мм.гггг, когда был '
                                                     f'осуществлён бракоразводный процесс и комментарии '
                                                     f'по деталям этого развода (А затем по следующим, если такие '
                                                     f'были).'
                                                     f' Информацию необходимо указывать в '
                                                     f'следующем формате: Развод 1-й дата, указанная в документе '
                                                     f'чч.мм.гггг. И комментарии по событиям этого развода.')
        await Fsm_storage.divorce.set()


async def nohandler9(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['divorce'] = 'Разводов не было'
        await state.finish()
        await bot.send_message(message.from_user.id, f'Это очень хорошо!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Есть/были ли у вас дети?', reply_markup=yes11)


async def divorce_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['divorce'] += f', {message.text}'
        except Exception as e:
            data['divorce'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes10)


async def yeshandler10(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Дополните ответным сообщением.')
        await Fsm_storage.divorce.set()


async def nohandler10(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Далее, у вас есть/были дети?', reply_markup=yes11)


async def yeshandler11(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Напишите пожалуйста полную Дату рождения Вашего ребёнка и '
                                                     f'комментарии относительно Рождения ребёнка в следующем формате: '
                                                     f'Ребёнок 1-й дата рождения чч.мм.гггг, '
                                                     f'комментарий относительно Рождения.')
        await Fsm_storage.children.set()


async def nohandler11(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['children'] = 'Детей нет и не было'
        await state.finish()
        await bot.send_message(message.from_user.id, f'Принято, идем дальше!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'В вашей жизни случались переезды в другой город/страну?',
                               reply_markup=yes13)


async def children_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['children'] += f', {message.text}'
        except Exception as e:
            data['children'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes12)


async def yeshandler12(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Дополните ответным сообщением.')
        await Fsm_storage.children.set()


async def nohandler12(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'В вашей жизни случались переезды в другой город/страну?',
                               reply_markup=yes13)


async def yeshandler13(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию о переездах, случавшихся в '
                                                     f'Вашей жизни (из города в город, из страны в страну), в '
                                                     f'следующем формате: Переезд из ____ в ____   дата переезда '
                                                     f'чч.мм.гггг, комментарий относительно этого переезда '
                                                     f'(вынужденный, по работе, в связи с какими обстоятельствами, '
                                                     f'как сказался на Вашей жизни).')
        await Fsm_storage.relocation.set()


async def nohandler13(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        data['relocation'] = 'Переездов не было'
        await state.finish()
        await bot.send_message(message.from_user.id, f'Принято, идем дальше!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию обо всех Ваших местах работы '
                                                     f'(приём/увольнение) в следующем формате: Работа 1-я дата приёма '
                                                     f'(чч.мм.гггг), дата увольнения (чч.мм.гггг), комментарий '
                                                     f'относительно этого места работы.')
        await Fsm_storage.job.set()


async def relocation_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['relocation'] += f', {message.text}'
        except Exception as e:
            data['relocation'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes14)


async def yeshandler14(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Дополните ответным сообщением.')
        await Fsm_storage.relocation.set()


async def nohandler14(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично!')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию обо всех Ваших местах работы '
                                                     f'(приём/увольнение) в следующем формате: Работа 1-я дата приёма '
                                                     f'(чч.мм.гггг), дата увольнения (чч.мм.гггг), комментарий '
                                                     f'относительно этого места работы.')
        await Fsm_storage.job.set()


async def job_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['job'] += f', {message.text}'
        except Exception as e:
            data['job'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить про работу, например еще одно место?',
                               reply_markup=yes15)


async def yeshandler15(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.job.set()


async def nohandler15(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично!')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               f'Укажите пожалуйста были ли в Вашей жизни события, в результате которых Вы '
                               f'получали награды, достигали побед в различных сферах, либо просто значимые достижения '
                               f'(необходимо написать все самые выдающееся события) в следующем формате: '
                               f'\n\nСобытие 1-е дата получения чч.мм.гггг, комментарий относительно этого достижения.')
        await Fsm_storage.events.set()


async def events_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['events'] += f', {message.text}'
        except Exception as e:
            data['events'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               'Хотите что-то добавить? возможно еще достижения или важные события?',
                               reply_markup=yes16)


async def yeshandler16(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.events.set()


async def nohandler16(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично. Перейдем к следующему вопросу.')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               f'Укажите пожалуйста информацию относительно смерти родственников/близких людей, '
                               f'в следующем формате: \n\nСобытие 1-е дата происшедшего чч.мм.гггг, комментарий '
                               f'относительно этой даты.')
        await Fsm_storage.death.set()


async def death_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['death'] += f', {message.text}'
        except Exception as e:
            data['death'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               'Хотите добавить еще событие?',
                               reply_markup=yes17)


async def yeshandler17(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.death.set()


async def nohandler17(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Здорово, двигаемся дальше.')
        time.sleep(1)
        await bot.send_message(message.from_user.id, f'Укажите пожалуйста информацию относительно травм/операций в '
                                                     f'Вашей жизни. \n\nСобытие 1-е дата случившегося чч.мм.гггг, '
                                                     f'комментарий относительно этого происшествия.')
        await Fsm_storage.operations.set()


async def operations_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['operations'] += f', {message.text}'
        except Exception as e:
            data['operations'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes18)


async def yeshandler18(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.operations.set()


async def nohandler18(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Спасибо, далее, следующий вопрос.')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               f'Укажите пожалуйста информацию о покупке движимого и недвижимого имущества '
                               f'(автомобиль, квартира и т.д.), в следующем формате: \n\nСобытие 1-е дата приобретения '
                               f'чч.мм.гггг, комментарий относительно этой покупки.')
        await Fsm_storage.property.set()


async def property_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['property'] += f', {message.text}'
        except Exception as e:
            data['property'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes19)


async def yeshandler19(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.property.set()

async def nohandler19(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Отлично. Перейдем к следующему вопросу.')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               f'Если в Вашей жизни происходило иное значимое событие, которое не упоминалось в '
                               f'вышеуказанном списке вопросов, просьбы указать его в следующем формате: \n\nСобытие '
                               f'1-е и его дата чч.мм.гггг, комментарий относительно происшедшего. Если добавить '
                               f'больше нечего, так и напишите.')
        await Fsm_storage.other.set()

async def other_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        try:
            data['other'] += f', {message.text}'
        except Exception as e:
            data['other'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Хотите что-то добавить?', reply_markup=yes20)

async def yeshandler20(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id, f'Хорошо, тогда дополните ответным сообщением.')
        await Fsm_storage.other.set()

async def nohandler20(message: types.CallbackQuery, state: FSMContext):
    async with state.proxy() as data:
        await bot.send_message(message.from_user.id,
                               f'Что Вас интересует в разборе Натальной карты/Прогностике? На чём следует '
                               f'акцентировать внимание? Так же можете просто перечислить сферы жизни, которые '
                               f'хочется обсудить во время консультации (личная жизнь, финансы, работа и пр.).')
        await Fsm_storage.wishes.set()

async def wishes_check(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['wishes'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id, 'Укажите пожалуйста Ваши контактные данные – номер телефона '
                                                     'и адрес электронной почты. Чтобы я могла с вами связаться')

        await Fsm_storage.contacts.set()

async def final(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['contacts'] = message.text
        await state.finish()
        await bot.send_message(message.from_user.id,
                               'Благодарю Вас за время, которое Вы уделили, чтобы ответить на все вышеперечисленные '
                               'вопросы. Хорошего вам дня! И До встречи!')
        time.sleep(1)
        await bot.send_message(message.from_user.id,
                               'Мои контакты для связи\nАлла Бабенко\nFreyaforseti0609@mail.ru '
                               '\nFreyaforseti0609@gmail.com\n+79998482603')
        time.sleep(2)
        await bot.send_message(message.from_user.id,
                               'Можете сохранить номер телефона')
        qr = open("static/qr.jpg", 'rb')
        await bot.send_photo(message.from_user.id, photo=qr)


        # login = f'@{message.from_user.username}'
        data_list = [data['fio'], f"{data['day_of_birth']}.{data['month_of_birth']}.{data['year_of_birth']}",
                     data['birth_time'], data['source_of_info'], data['circumstances'], data['marriage'],
                     data['divorce'], data['children'], data['relocation'], data['job'], data['events'], data['death'],
                     data['operations'], data['property'], data['other'], data['wishes'], data['contacts']]

        list_of_points = ['ФИО', 'Дата рождения', 'Время рождения', 'Источник информации о времени рождения',
                          'Обстоятельства рождения', 'Брак', 'Развод', 'Дети', 'Переезд', 'Работа',
                          'Важные события в жизни', 'Смерть родственников', 'Травмы и операции',
                          'Частная собственность', 'Прочие события', 'Уточнение по Консультации', 'Контакты для связи']
        mydoc = docx.Document()

        mydoc.add_heading(f"Анкета клиента {data['fio']}", 1)
        mydoc.add_paragraph('')
        for i in range(len(list_of_points)):
            p = mydoc.add_paragraph('')
            p.add_run(f'{list_of_points[i]}: ').bold = True
            p.add_run(f'{data_list[i]}')

        mydoc.save(f"{data['fio']}.docx")
        file = open(f"{data['fio']}.docx", 'rb')
        await bot.send_document(363700041, document=file)
        try:
            await bot.send_message(363700041, f'Пользователь: @{message.from_user.username}')
        except:
            await bot.send_message(363700041, f'У пользователя нет username. Необходимо проверить "Контакты для связи" '
                                              f'в файле выше.')
        os.remove(f"{data['fio']}.docx")
    await state.finish()
    await state.reset_state()


async def stats(message: types.Message):
    print(message)
    print(message.from_user)
    await bot.send_message(363700041, f'{message.from_user.id}')
    await bot.send_message(363700041, f'Пользователь: {message.from_user.username}, id: {message.from_user.id}')


async def lost(message: types.Message):
    await bot.send_message(message.from_user.id, 'Не понимаю Вас, нажмите --- /start')


def register_handlers_client(dp: Dispatcher):
    dp.register_message_handler(main, commands=['start', 'help'])
    dp.register_message_handler(gotov, commands=['Готов'])
    dp.register_message_handler(negotov, commands=['Не_Готов'])
    dp.register_message_handler(fio_check, state=Fsm_storage.fio)
    dp.register_callback_query_handler(yeshandler1, text='yes1')
    dp.register_callback_query_handler(nohandler1, text='no1')
    dp.register_message_handler(year_of_birth_check, state=Fsm_storage.year_of_birth)
    dp.register_message_handler(month_of_birth_check, state=Fsm_storage.month_of_birth)
    dp.register_message_handler(day_of_birth_check, state=Fsm_storage.day_of_birth)
    dp.register_message_handler(birth_time_check, state=Fsm_storage.birth_time)
    dp.register_message_handler(place_of_birth_check, state=Fsm_storage.place_of_birth)
    dp.register_message_handler(source_of_info_check, state=Fsm_storage.source_of_info)
    dp.register_message_handler(circumstances_check, state=Fsm_storage.circumstances)
    dp.register_callback_query_handler(yeshandler6, text='yes6')
    dp.register_callback_query_handler(nohandler6, text='no6')
    dp.register_callback_query_handler(yeshandler7, text='yes7')
    dp.register_callback_query_handler(nohandler7, text='no7')
    dp.register_message_handler(marriage_check, state=Fsm_storage.marriage)
    dp.register_callback_query_handler(yeshandler8, text='yes8')
    dp.register_callback_query_handler(nohandler8, text='no8')
    dp.register_callback_query_handler(yeshandler9, text='yes9')
    dp.register_callback_query_handler(nohandler9, text='no9')
    dp.register_message_handler(divorce_check, state=Fsm_storage.divorce)
    dp.register_callback_query_handler(yeshandler10, text='yes10')
    dp.register_callback_query_handler(nohandler10, text='no10')
    dp.register_callback_query_handler(yeshandler11, text='yes11')
    dp.register_callback_query_handler(nohandler11, text='no11')
    dp.register_message_handler(children_check, state=Fsm_storage.children)
    dp.register_callback_query_handler(yeshandler12, text='yes12')
    dp.register_callback_query_handler(nohandler12, text='no12')
    dp.register_callback_query_handler(yeshandler13, text='yes13')
    dp.register_callback_query_handler(nohandler13, text='no13')
    dp.register_message_handler(relocation_check, state=Fsm_storage.relocation)
    dp.register_callback_query_handler(yeshandler14, text='yes14')
    dp.register_callback_query_handler(nohandler14, text='no14')
    dp.register_message_handler(job_check, state=Fsm_storage.job)
    dp.register_callback_query_handler(yeshandler15, text='yes15')
    dp.register_callback_query_handler(nohandler15, text='no15')
    dp.register_message_handler(events_check, state=Fsm_storage.events)
    dp.register_callback_query_handler(yeshandler16, text='yes16')
    dp.register_callback_query_handler(nohandler16, text='no16')
    dp.register_message_handler(death_check, state=Fsm_storage.death)
    dp.register_callback_query_handler(yeshandler17, text='yes17')
    dp.register_callback_query_handler(nohandler17, text='no17')
    dp.register_message_handler(operations_check, state=Fsm_storage.operations)
    dp.register_callback_query_handler(yeshandler18, text='yes18')
    dp.register_callback_query_handler(nohandler18, text='no18')
    dp.register_message_handler(property_check, state=Fsm_storage.property)
    dp.register_callback_query_handler(yeshandler19, text='yes19')
    dp.register_callback_query_handler(nohandler19, text='no19')
    dp.register_message_handler(other_check, state=Fsm_storage.other)
    dp.register_callback_query_handler(yeshandler20, text='yes20')
    dp.register_callback_query_handler(nohandler20, text='no20')
    dp.register_message_handler(wishes_check, state=Fsm_storage.wishes)
    dp.register_message_handler(final, state=Fsm_storage.contacts)
    dp.register_message_handler(stats, commands=['stats'])
    dp.register_message_handler(lost)
